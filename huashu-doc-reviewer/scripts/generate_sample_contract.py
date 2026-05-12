#!/usr/bin/env python3
"""生成演示用的示例合同 .docx。

脱敏合同，几处埋好了风险点，用于演示批注+修订。

用法：
    python3 generate_sample_contract.py --output examples/sample_contract.docx
"""

import argparse
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


CONTRACT_PARAGRAPHS = [
    ("《商业合作协议》", "Title"),
    ("", "Normal"),
    ("甲方：XX科技有限公司", "Normal"),
    ("乙方：花叔（自媒体博主）", "Normal"),
    ("签订日期：2026年4月17日", "Normal"),
    ("", "Normal"),
    ("一、合作内容", "Heading 1"),
    ("根据本合同，乙方需在其公众号、B站、小红书等平台，为甲方产品进行推广。具体推广形式包括但不限于文章、视频、图文笔记等。甲方应提供必要的产品资料和素材。", "Normal"),
    ("", "Normal"),
    ("二、合作期限", "Heading 1"),
    ("本合同自签订之日起生效，有效期为壹年。如遇不可抗力，双方可协商变更或终止合同。", "Normal"),
    ("", "Normal"),
    ("三、费用与结算", "Heading 1"),
    ("甲方应于内容发布后30日内，向乙方支付合作费用共计10万元整。如甲方逾期付款，乙方可要求赔偿。", "Normal"),
    ("乙方开具增值税专用发票后，甲方方可办理付款手续。", "Normal"),
    ("", "Normal"),
    ("四、内容审核", "Heading 1"),
    ("乙方所创作的所有推广内容，须经甲方审核通过后方可发布。甲方对内容的修改意见，乙方应予以配合。若内容发布后，甲方发现存在问题，乙方应无条件配合删除或修改。", "Normal"),
    ("", "Normal"),
    ("五、违约责任", "Heading 1"),
    ("任何一方违反本合同约定，应向守约方支付违约金。违约金数额由双方协商确定。因违约造成的损失，违约方应予以赔偿。", "Normal"),
    ("", "Normal"),
    ("六、知识产权", "Heading 1"),
    ("乙方创作的所有内容的知识产权归甲方所有。乙方不得将相关内容用于其他商业用途。", "Normal"),
    ("", "Normal"),
    ("七、保密条款", "Heading 1"),
    ("双方对本合同内容及合作过程中知悉的对方商业信息负有保密义务。", "Normal"),
    ("", "Normal"),
    ("八、争议解决", "Heading 1"),
    ("本合同项下发生争议的，双方应友好协商解决。协商不成的，任何一方均可向甲方所在地人民法院提起诉讼。", "Normal"),
    ("", "Normal"),
    ("九、其他事项", "Heading 1"),
    ("本合同一式两份，双方各执一份，具有同等法律效力。本合同未尽事宜，双方可另行签订补充协议。", "Normal"),
    ("", "Normal"),
    ("甲方（盖章）：_____________        乙方（签字）：_____________", "Normal"),
    ("日期：_______________              日期：_______________", "Normal"),
]


def build(output_path: str) -> None:
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "SimSun"
    font.size = Pt(11)

    for text, style_name in CONTRACT_PARAGRAPHS:
        if style_name == "Title":
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0] if p.runs else None
            if run:
                run.font.size = Pt(18)
                run.font.bold = True
        elif style_name == "Heading 1":
            p = doc.add_heading(text, level=1)
            run = p.runs[0] if p.runs else None
            if run:
                run.font.size = Pt(14)
        else:
            doc.add_paragraph(text)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Wrote: {output_path}", file=sys.stderr)


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate sample contract .docx")
    parser.add_argument("--output", "-o", default="sample_contract.docx",
                        help="Output .docx path")
    args = parser.parse_args()

    try:
        build(args.output)
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
