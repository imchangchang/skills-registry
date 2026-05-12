#!/usr/bin/env python3
"""读取 docx 并输出结构化 JSON，给模型做"判断"。

模型只看到 {paragraphs: [...], tables: [...]}，
不碰 XML、不碰格式、不碰原文件。

用法：
    python3 read_docx.py <file.docx>
    python3 read_docx.py <file.docx> --output doc.json
    python3 read_docx.py <file.docx> --for-model   # 给模型的精简版（无 full_text）
    python3 read_docx.py <file.docx> --max-paragraph-chars 2000  # 超长段落截断

依赖：python-docx（pip install python-docx）
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def read_docx(path: str, max_paragraph_chars: int = 0) -> dict:
    """读取 docx 文件，返回结构化字典。

    max_paragraph_chars > 0 时，超长段落会被截断（在提示给模型时避免 context 爆炸）。
    注入批注时用的是完整段落文本，不受截断影响。
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")
    if p.suffix.lower() != ".docx":
        raise ValueError(f"Expected .docx, got: {p.suffix}")

    doc = Document(str(p))

    paragraphs = []
    total_chars = 0
    for idx, para in enumerate(doc.paragraphs):
        text = para.text or ""
        total_chars += len(text)
        display_text = text
        if max_paragraph_chars and len(text) > max_paragraph_chars:
            display_text = text[:max_paragraph_chars] + f"...[truncated {len(text) - max_paragraph_chars} chars]"
        paragraphs.append({
            "id": idx,
            "text": display_text,
            "full_text": text,  # 完整文本（代码注入时用这个定位）
            "style": para.style.name if para.style else "Normal",
        })

    tables = []
    for t_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append({
            "id": t_idx,
            "rows": rows,
        })

    return {
        "source": str(p.resolve()),
        "paragraphs": paragraphs,
        "tables": tables,
        "stats": {
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "total_chars": total_chars,
        },
    }


def to_model_prompt(data: dict) -> str:
    """生成给模型看的精简版——剥掉 full_text，不含空段落。"""
    slim_paras = [
        {"id": p["id"], "text": p["text"], "style": p["style"]}
        for p in data["paragraphs"]
        if (p["text"] or "").strip()
    ]
    slim = {
        "paragraphs": slim_paras,
        "tables": data["tables"],
        "stats": data["stats"],
    }
    return json.dumps(slim, ensure_ascii=False, indent=2)


def main() -> int:
    parser = argparse.ArgumentParser(description="Read .docx and output structured JSON")
    parser.add_argument("input", help="Input .docx file")
    parser.add_argument("--output", "-o", help="Output JSON file (default: stdout)")
    parser.add_argument("--max-paragraph-chars", type=int, default=0,
                        help="Truncate paragraphs longer than this (0 = no truncation)")
    parser.add_argument("--for-model", action="store_true",
                        help="Output slim JSON for model prompt (no full_text field)")
    args = parser.parse_args()

    try:
        data = read_docx(args.input, max_paragraph_chars=args.max_paragraph_chars)
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        return 1

    if args.for_model:
        out = to_model_prompt(data)
    else:
        out = json.dumps(data, ensure_ascii=False, indent=2)

    if args.output:
        Path(args.output).write_text(out, encoding="utf-8")
        print(f"Wrote: {args.output}", file=sys.stderr)
    else:
        print(out)

    return 0


if __name__ == "__main__":
    sys.exit(main())
